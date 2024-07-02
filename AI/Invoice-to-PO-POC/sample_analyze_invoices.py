"""
This code sample shows Prebuilt Invoice operations with the Azure Form Recognizer client library. 
The async versions of the samples require Python 3.6 or later.

To learn more, please visit the documentation - Quickstart: Form Recognizer Python client library SDKs
https://learn.microsoft.com/azure/applied-ai-services/form-recognizer/quickstarts/get-started-v3-sdk-rest-api?view=doc-intel-3.1.0&pivots=programming-language-python
"""
import docx
import os

from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeResult

from azure.core.credentials import AzureKeyCredential
from azure.ai.formrecognizer import DocumentAnalysisClient
from docxtpl import DocxTemplate

"""
Remember to remove the key from your code when you're done, and never post it publicly. For production, use
secure methods to store and access your credentials. For more information, see 
https://docs.microsoft.com/en-us/azure/cognitive-services/cognitive-services-security?tabs=command-line%2Ccsharp#environment-variables-and-application-configuration
"""
endpoint = "https://inv-to-po-poc.cognitiveservices.azure.com/"
key = ""

# sample document
formUrl = "https://raw.githubusercontent.com/Azure-Samples/cognitive-services-REST-api-samples/master/curl/form-recognizer/sample-invoice.pdf"

document_analysis_client = DocumentAnalysisClient(
    endpoint=endpoint, credential=AzureKeyCredential(key)
)
    
poller = document_analysis_client.begin_analyze_document_from_url("prebuilt-invoice", formUrl)
invoices = poller.result()

for idx, invoice in enumerate(invoices.documents):
    doc = DocxTemplate("purchase-order.docx")
    print("--------Recognizing invoice #{}--------".format(idx + 1))
    vendor_name = invoice.fields.get("VendorName")
    if vendor_name:
        print(
            "Vendor Name: {} has confidence: {}".format(
                vendor_name.value, vendor_name.confidence
            )
        )
        context = { 'CompanyName' : f"{vendor_name.value}" }
        doc.render(context)
    vendor_address = invoice.fields.get("VendorAddress")
    if vendor_address:
        print(
            "Vendor Address: {} has confidence: {}".format(
                vendor_address.value, vendor_address.confidence
            )
        )
        context1 = { 'Address' : f"{vendor_address.value.house_number}", 'CompanyName' : f"{vendor_name.value}" }
        doc.render(context1)
    vendor_address_recipient = invoice.fields.get("VendorAddressRecipient")
    if vendor_address_recipient:
        print(
            "Vendor Address Recipient: {} has confidence: {}".format(
                vendor_address_recipient.value, vendor_address_recipient.confidence
            )
        )
    customer_name = invoice.fields.get("CustomerName")
    if customer_name:
        print(
            "Customer Name: {} has confidence: {}".format(
                customer_name.value, customer_name.confidence
            )
        )
    customer_id = invoice.fields.get("CustomerId")
    if customer_id:
        print(
            "Customer Id: {} has confidence: {}".format(
                customer_id.value, customer_id.confidence
            )
        )
    customer_address = invoice.fields.get("CustomerAddress")
    if customer_address:
        print(
            "Customer Address: {} has confidence: {}".format(
                customer_address.value, customer_address.confidence
            )
        )
    customer_address_recipient = invoice.fields.get("CustomerAddressRecipient")
    if customer_address_recipient:
        print(
            "Customer Address Recipient: {} has confidence: {}".format(
                customer_address_recipient.value,
                customer_address_recipient.confidence,
            )
        )
    invoice_id = invoice.fields.get("InvoiceId")
    if invoice_id:
        print(
            "Invoice Id: {} has confidence: {}".format(
                invoice_id.value, invoice_id.confidence
            )
        )
    invoice_date = invoice.fields.get("InvoiceDate")
    if invoice_date:
        print(
            "Invoice Date: {} has confidence: {}".format(
                invoice_date.value, invoice_date.confidence
            )
        )
    invoice_total = invoice.fields.get("InvoiceTotal")
    if invoice_total:
        print(
            "Invoice Total: {} has confidence: {}".format(
                invoice_total.value, invoice_total.confidence
            )
        )
    due_date = invoice.fields.get("DueDate")
    if due_date:
        print(
            "Due Date: {} has confidence: {}".format(
                due_date.value, due_date.confidence
            )
        )
    purchase_order = invoice.fields.get("PurchaseOrder")
    if purchase_order:
        print(
            "Purchase Order: {} has confidence: {}".format(
                purchase_order.value, purchase_order.confidence
            )
        )
    billing_address = invoice.fields.get("BillingAddress")
    if billing_address:
        print(
            "Billing Address: {} has confidence: {}".format(
                billing_address.value, billing_address.confidence
            )
        )
    billing_address_recipient = invoice.fields.get("BillingAddressRecipient")
    if billing_address_recipient:
        print(
            "Billing Address Recipient: {} has confidence: {}".format(
                billing_address_recipient.value,
                billing_address_recipient.confidence,
            )
        )
    shipping_address = invoice.fields.get("ShippingAddress")
    if shipping_address:
        print(
            "Shipping Address: {} has confidence: {}".format(
                shipping_address.value, shipping_address.confidence
            )
        )
    shipping_address_recipient = invoice.fields.get("ShippingAddressRecipient")
    if shipping_address_recipient:
        print(
            "Shipping Address Recipient: {} has confidence: {}".format(
                shipping_address_recipient.value,
                shipping_address_recipient.confidence,
            )
        )
    print("Invoice items:")
    for idx, item in enumerate(invoice.fields.get("Items").value):
        print("...Item #{}".format(idx + 1))
        item_description = item.value.get("Description")
        if item_description:
            print(
                "......Description: {} has confidence: {}".format(
                    item_description.value, item_description.confidence
                )
            )
        item_quantity = item.value.get("Quantity")
        if item_quantity:
            print(
                "......Quantity: {} has confidence: {}".format(
                    item_quantity.value, item_quantity.confidence
                )
            )
        unit = item.value.get("Unit")
        if unit:
            print(
                "......Unit: {} has confidence: {}".format(
                    unit.value, unit.confidence
                )
            )
        unit_price = item.value.get("UnitPrice")
        if unit_price:
            print(
                "......Unit Price: {} has confidence: {}".format(
                    unit_price.value, unit_price.confidence
                )
            )
        product_code = item.value.get("ProductCode")
        if product_code:
            print(
                "......Product Code: {} has confidence: {}".format(
                    product_code.value, product_code.confidence
                )
            )
        item_date = item.value.get("Date")
        if item_date:
            print(
                "......Date: {} has confidence: {}".format(
                    item_date.value, item_date.confidence
                )
            )
        tax = item.value.get("Tax")
        if tax:
            print(
                "......Tax: {} has confidence: {}".format(tax.value, tax.confidence)
            )
        amount = item.value.get("Amount")
        if amount:
            print(
                "......Amount: {} has confidence: {}".format(
                    amount.value, amount.confidence
                )
            )
    subtotal = invoice.fields.get("SubTotal")
    if subtotal:
        print(
            "Subtotal: {} has confidence: {}".format(
                subtotal.value, subtotal.confidence
            )
        )
    total_tax = invoice.fields.get("TotalTax")
    if total_tax:
        print(
            "Total Tax: {} has confidence: {}".format(
                total_tax.value, total_tax.confidence
            )
        )
    previous_unpaid_balance = invoice.fields.get("PreviousUnpaidBalance")
    if previous_unpaid_balance:
        print(
            "Previous Unpaid Balance: {} has confidence: {}".format(
                previous_unpaid_balance.value, previous_unpaid_balance.confidence
            )
        )
    amount_due = invoice.fields.get("AmountDue")
    if amount_due:
        print(
            "Amount Due: {} has confidence: {}".format(
                amount_due.value, amount_due.confidence
            )
        )
    service_start_date = invoice.fields.get("ServiceStartDate")
    if service_start_date:
        print(
            "Service Start Date: {} has confidence: {}".format(
                service_start_date.value, service_start_date.confidence
            )
        )
    service_end_date = invoice.fields.get("ServiceEndDate")
    if service_end_date:
        print(
            "Service End Date: {} has confidence: {}".format(
                service_end_date.value, service_end_date.confidence
            )
        )
    service_address = invoice.fields.get("ServiceAddress")
    if service_address:
        print(
            "Service Address: {} has confidence: {}".format(
                service_address.value, service_address.confidence
            )
        )
    service_address_recipient = invoice.fields.get("ServiceAddressRecipient")
    if service_address_recipient:
        print(
            "Service Address Recipient: {} has confidence: {}".format(
                service_address_recipient.value,
                service_address_recipient.confidence,
            )
        )
    remittance_address = invoice.fields.get("RemittanceAddress")
    if remittance_address:
        print(
            "Remittance Address: {} has confidence: {}".format(
                remittance_address.value, remittance_address.confidence
            )
        )
    remittance_address_recipient = invoice.fields.get("RemittanceAddressRecipient")
    if remittance_address_recipient:
        print(
            "Remittance Address Recipient: {} has confidence: {}".format(
                remittance_address_recipient.value,
                remittance_address_recipient.confidence,
            )
        )
    print("----------------------------------------")
    doc.save("generated_doc.docx")

for idx, invoice in enumerate(invoices.documents):
    doc = docx.Document()
    print("--------Recognizing invoice for doc #{}--------".format(idx + 1))
    vendor_name = invoice.fields.get("VendorName")
    if vendor_name:
        doc = docx.Document()
        p = doc.add_paragraph()
        run = p.add_run(f"Vendor = {vendor_name.value}")
        doc.save("docdemo1.docx") 
        
# Create a document
doc = docx.Document()

# Add a paragraph to the document
p = doc.add_paragraph()

# Add a run to the paragraph
run = p.add_run("python-docx")

# Add more text to the same paragraph
run = p.add_run(" Tutorial")

# Add another paragraph (left blank for an empty line)
doc.add_paragraph()

# Add another paragraph
p = doc.add_paragraph()

# Add a run and format it
run = p.add_run("This is my first python-docx tutorial!")


# Save the document
doc.save("docdemo.docx") 